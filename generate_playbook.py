"""
Generate UI Playbook for Stratos BEP.
Takes screenshots of every page for all 3 roles and assembles a Word document.

Usage:
    python generate_playbook.py [--url http://178.104.202.193]
"""
import argparse
import os
import time

from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCREENSHOT_DIR = 'docs/playbook_screenshots'
OUTPUT_PATH = 'docs/PLAYBOOK.docx'

USERS = {
    'admin': {'password': 'admin123', 'role': 'ADMIN'},
    'analyst': {'password': 'analyst123', 'role': 'ANALYST'},
    'viewer': {'password': 'viewer123', 'role': 'VIEWER'},
}

# Define all screenshots to take per role
# Format: (filename, url_path, description, wait_seconds)
ADMIN_PAGES = [
    ('01_login', '/accounts/login/', 'The login page provides secure authentication for all users. Enter your username and password to access the Stratos BEP dashboard.', 1),
    ('02_dashboard', '/', 'The Admin dashboard displays real-time email security statistics: total emails processed, clean/suspicious/malicious counts, recent alerts, and threat intelligence overview. The sidebar shows all navigation sections including the Admin panel.', 2),
    ('03_email_list', '/emails/', 'The email list page shows all processed emails with their verdict (CLEAN, SUSPICIOUS, MALICIOUS), score, status, and timestamp. Filters for verdict, status, sender address, and date range are available at the top.', 2),
    ('04_email_list_filtered', '/emails/?verdict=MALICIOUS', 'The email list filtered by MALICIOUS verdict, showing only emails that scored 70 or above. This helps analysts quickly focus on the most critical threats.', 2),
    ('05_email_detail_malicious', None, 'The email detail page for a malicious email showing the score breakdown across all pipeline stages: Preprocessor (SPF/DKIM/DMARC checks), Keywords, URLs, Attachments, and Chain Analysis. The large score badge and verdict are prominently displayed.', 2),
    ('06_email_detail_raw', None, 'The Raw Analysis tab shows the complete JSON output from the analysis pipeline. This is only visible to ADMIN and ANALYST roles, providing full transparency into how the verdict was determined.', 1),
    ('07_quarantine', '/quarantine/', 'The quarantine management page lists all emails pending review. ADMIN users see Release, Block, and Delete action buttons for each entry. The status filter allows narrowing results.', 2),
    ('08_iocs', '/iocs/', 'The Indicators of Compromise (IOC) page displays all extracted threat indicators: malicious domains, URLs, IP addresses, and file hashes. Each IOC is linked to the source email and shows severity level.', 2),
    ('09_threat_intel', '/threat-intel/', 'The Threat Intelligence page shows feed statistics (MalwareBazaar hashes, URLhaus domains, malicious IPs, YARA rules), sync status, and whitelist/blacklist management. ADMIN users can add/remove entries and trigger manual syncs.', 2),
    ('10_reports', '/reports/', 'The Reports page provides data export capabilities: Email Summary (CSV), IOC List (CSV), and TI Stats (JSON). Report history and scheduled report configuration are also shown.', 2),
    ('11_users', '/users/', 'User Management allows ADMIN users to create new accounts, change user roles (ADMIN/ANALYST/VIEWER), and activate/deactivate accounts. Self-demotion and self-deactivation are prevented.', 2),
    ('12_settings', '/settings/', 'System Settings page (ADMIN only) provides Gmail OAuth integration, API key management (VirusTotal, AbuseIPDB), detection threshold configuration, and fetch interval settings. API keys are encrypted at rest.', 2),
]

ANALYST_PAGES = [
    ('13_analyst_dashboard', '/', 'The Analyst dashboard shows the same statistics as ADMIN but without the Admin section in the sidebar. Analysts cannot access User Management, Settings, or Django Admin.', 2),
    ('14_analyst_email_detail', None, 'Analysts can view the full email detail including the Raw Analysis JSON tab, giving them visibility into scoring details for investigation purposes.', 2),
    ('15_analyst_quarantine', '/quarantine/', 'Analysts have full quarantine action capabilities: Release (deliver email), Block (block sender and add to blacklist), and Delete (permanently remove). These actions are identical to ADMIN.', 2),
    ('16_analyst_reports', '/reports/', 'Analysts can access Email Summary and IOC List exports but cannot export TI Stats (ADMIN only) or manage scheduled reports.', 2),
    ('17_analyst_threat_intel', '/threat-intel/', 'Analysts can view Threat Intelligence statistics and feed status but cannot add/remove whitelist or blacklist entries, and cannot trigger manual syncs.', 2),
]

VIEWER_PAGES = [
    ('18_viewer_dashboard', '/', 'The Viewer dashboard provides read-only access to email security statistics. The sidebar shows Monitor, Security, and Reports sections only — no Admin panel.', 2),
    ('19_viewer_email_detail', None, 'Viewers can see email details but the Raw Analysis tab is hidden. They can view the basic score breakdown but not the full pipeline JSON output.', 2),
    ('20_viewer_quarantine', '/quarantine/', 'Viewers can see the quarantine list but have no action buttons. They cannot release, block, or delete quarantined emails — this is read-only access.', 2),
    ('21_viewer_reports', '/reports/', 'Viewers can see the Reports page and report history but have no export buttons. Data export requires at least ANALYST role.', 2),
    ('22_viewer_threat_intel', '/threat-intel/', 'Viewers can see Threat Intelligence statistics and feed status in read-only mode. No management buttons (sync, whitelist/blacklist add/remove) are visible.', 2),
]


def login(page, base_url, username, password):
    """Log in as a specific user."""
    page.goto(f'{base_url}/accounts/login/')
    page.fill('input[name="username"]', username)
    page.fill('input[name="password"]', password)
    page.click('button[type="submit"]')
    page.wait_for_load_state('networkidle')
    time.sleep(1)


def logout(page, base_url):
    """Log out."""
    page.goto(f'{base_url}/accounts/logout/')
    page.wait_for_load_state('networkidle')


def take_screenshot(page, filename):
    """Take a full-page screenshot."""
    path = os.path.join(SCREENSHOT_DIR, f'{filename}.png')
    page.screenshot(path=path, full_page=True)
    print(f'  Screenshot: {path}')
    return path


def find_malicious_email_link(page):
    """Find and return the URL of the first malicious email detail page."""
    links = page.query_selector_all('table tbody tr')
    for row in links:
        verdict_cell = row.query_selector('td:nth-child(4)')
        if verdict_cell and 'MALICIOUS' in (verdict_cell.inner_text() or ''):
            link = row.query_selector('td:nth-child(3) a')
            if link:
                return link.get_attribute('href')
    return None


def capture_admin_screenshots(page, base_url):
    """Capture all ADMIN role screenshots."""
    print('\n=== ADMIN Role ===')
    login(page, base_url, 'admin', 'admin123')
    screenshots = []

    for filename, url_path, description, wait in ADMIN_PAGES:
        if url_path:
            page.goto(f'{base_url}{url_path}')
            page.wait_for_load_state('networkidle')
            time.sleep(wait)

        if filename == '01_login':
            # Already on login, need to logout first for this screenshot
            logout(page, base_url)
            page.goto(f'{base_url}/accounts/login/')
            page.wait_for_load_state('networkidle')
            time.sleep(1)
            path = take_screenshot(page, filename)
            screenshots.append((path, description))
            login(page, base_url, 'admin', 'admin123')
            continue

        if filename == '05_email_detail_malicious':
            # Navigate to email list first, find a malicious email
            page.goto(f'{base_url}/emails/?verdict=MALICIOUS')
            page.wait_for_load_state('networkidle')
            time.sleep(1)
            link = find_malicious_email_link(page)
            if link:
                page.goto(f'{base_url}{link}')
                page.wait_for_load_state('networkidle')
                time.sleep(wait)
            path = take_screenshot(page, filename)
            screenshots.append((path, description))
            continue

        if filename == '06_email_detail_raw':
            # Click the Raw Analysis tab on the current detail page
            raw_btn = page.query_selector('button:has-text("Raw Analysis")')
            if raw_btn:
                raw_btn.click()
                time.sleep(1)
            path = take_screenshot(page, filename)
            screenshots.append((path, description))
            continue

        path = take_screenshot(page, filename)
        screenshots.append((path, description))

    logout(page, base_url)
    return screenshots


def capture_analyst_screenshots(page, base_url):
    """Capture all ANALYST role screenshots."""
    print('\n=== ANALYST Role ===')
    login(page, base_url, 'analyst', 'analyst123')
    screenshots = []

    for filename, url_path, description, wait in ANALYST_PAGES:
        if url_path:
            page.goto(f'{base_url}{url_path}')
            page.wait_for_load_state('networkidle')
            time.sleep(wait)

        if filename == '14_analyst_email_detail':
            page.goto(f'{base_url}/emails/?verdict=MALICIOUS')
            page.wait_for_load_state('networkidle')
            time.sleep(1)
            link = find_malicious_email_link(page)
            if link:
                page.goto(f'{base_url}{link}')
                page.wait_for_load_state('networkidle')
                time.sleep(wait)
            path = take_screenshot(page, filename)
            screenshots.append((path, description))
            continue

        path = take_screenshot(page, filename)
        screenshots.append((path, description))

    logout(page, base_url)
    return screenshots


def capture_viewer_screenshots(page, base_url):
    """Capture all VIEWER role screenshots."""
    print('\n=== VIEWER Role ===')
    login(page, base_url, 'viewer', 'viewer123')
    screenshots = []

    for filename, url_path, description, wait in VIEWER_PAGES:
        if url_path:
            page.goto(f'{base_url}{url_path}')
            page.wait_for_load_state('networkidle')
            time.sleep(wait)

        if filename == '19_viewer_email_detail':
            page.goto(f'{base_url}/emails/')
            page.wait_for_load_state('networkidle')
            time.sleep(1)
            link = page.query_selector('table tbody tr:first-child td:nth-child(3) a')
            if link:
                href = link.get_attribute('href')
                page.goto(f'{base_url}{href}')
                page.wait_for_load_state('networkidle')
                time.sleep(wait)
            path = take_screenshot(page, filename)
            screenshots.append((path, description))
            continue

        path = take_screenshot(page, filename)
        screenshots.append((path, description))

    logout(page, base_url)
    return screenshots


def build_docx(admin_shots, analyst_shots, viewer_shots):
    """Build the Word document with all screenshots and descriptions."""
    doc = Document()

    # -- Styles --
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # -- Title Page --
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Stratos BEP')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(30, 58, 95)  # Navy
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('UI Playbook')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(37, 99, 235)  # Accent blue

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('Business Email Protection System\nUser Interface Guide for All Roles')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()
    roles_desc = doc.add_paragraph()
    roles_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = roles_desc.add_run('ADMIN  |  ANALYST  |  VIEWER')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(30, 58, 95)
    run.bold = True

    doc.add_page_break()

    # -- Table of Contents --
    toc = doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. ADMIN Role',
        '   1.1 Login Page',
        '   1.2 Dashboard',
        '   1.3 Email List',
        '   1.4 Email List (Filtered)',
        '   1.5 Email Detail (Malicious)',
        '   1.6 Email Detail (Raw Analysis)',
        '   1.7 Quarantine Management',
        '   1.8 IOC List',
        '   1.9 Threat Intelligence',
        '   1.10 Reports',
        '   1.11 User Management',
        '   1.12 System Settings',
        '2. ANALYST Role',
        '   2.1 Dashboard',
        '   2.2 Email Detail',
        '   2.3 Quarantine Management',
        '   2.4 Reports',
        '   2.5 Threat Intelligence',
        '3. VIEWER Role',
        '   3.1 Dashboard',
        '   3.2 Email Detail',
        '   3.3 Quarantine (Read-Only)',
        '   3.4 Reports (Read-Only)',
        '   3.5 Threat Intelligence (Read-Only)',
        '4. Role Permission Comparison',
        '5. System Settings — Step-by-Step Guide',
        '   5.1 Gmail Integration',
        '   5.2 Threat Intelligence API Keys',
        '   5.3 Detection Thresholds',
        '   5.4 Fetch Interval and TI Sync',
        '   5.5 Settings — DO and DO NOT',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.style.font.size = Pt(11)
        if not item.startswith('   '):
            p.runs[0].bold = True

    doc.add_page_break()

    # -- Helper function --
    def add_section(heading, shots, level=1):
        doc.add_heading(heading, level=level)
        for i, (img_path, description) in enumerate(shots):
            if os.path.exists(img_path):
                # Add image
                doc.add_picture(img_path, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Add caption
                caption_num = i + 1
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(f'Figure {caption_num}')
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(100, 116, 139)

                # Add description
                desc_para = doc.add_paragraph(description)
                desc_para.style.font.size = Pt(11)
                doc.add_paragraph()  # spacing
            else:
                doc.add_paragraph(f'[Screenshot not available: {img_path}]')

    # -- Role Sections --
    add_section('1. ADMIN Role — Full System Access', admin_shots)
    doc.add_page_break()
    add_section('2. ANALYST Role — Investigation & Response', analyst_shots)
    doc.add_page_break()
    add_section('3. VIEWER Role — Read-Only Monitoring', viewer_shots)

    # -- Role Comparison Table --
    doc.add_page_break()
    doc.add_heading('4. Role Permission Comparison', level=1)

    table = doc.add_table(rows=12, cols=4, style='Light Grid Accent 1')
    headers = ['Feature', 'ADMIN', 'ANALYST', 'VIEWER']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    permissions = [
        ('Dashboard', 'Full', 'Full', 'Full'),
        ('Email List & Detail', 'Full + Raw Analysis', 'Full + Raw Analysis', 'Basic (no raw)'),
        ('Quarantine Actions', 'Release / Block / Delete', 'Release / Block / Delete', 'View only'),
        ('IOC List', 'View + Export', 'View + Export', 'View only'),
        ('Threat Intel Management', 'Full (sync, whitelist, blacklist)', 'View only', 'View only'),
        ('Email Summary Export', 'Yes', 'Yes', 'No'),
        ('IOC Export', 'Yes', 'Yes', 'No'),
        ('TI Stats Export', 'Yes', 'No', 'No'),
        ('User Management', 'Full', 'No access', 'No access'),
        ('System Settings', 'Full', 'No access', 'No access'),
        ('Scheduled Reports', 'Manage', 'No access', 'No access'),
    ]
    for row_idx, (feature, admin, analyst, viewer) in enumerate(permissions, 1):
        table.rows[row_idx].cells[0].text = feature
        table.rows[row_idx].cells[1].text = admin
        table.rows[row_idx].cells[2].text = analyst
        table.rows[row_idx].cells[3].text = viewer

    # -- Section 5: Settings Page Guide --
    doc.add_page_break()
    doc.add_heading('5. System Settings — Step-by-Step Guide', level=1)

    p = doc.add_paragraph('The Settings page is accessible only to ADMIN users via the sidebar (Admin > Settings). '
                          'It controls Gmail integration, threat intelligence API keys, detection thresholds, and system behaviour.')
    doc.add_paragraph()

    # Add the settings screenshot
    settings_img = os.path.join(SCREENSHOT_DIR, '12_settings.png')
    if os.path.exists(settings_img):
        doc.add_picture(settings_img, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run('Settings Page Overview')
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(100, 116, 139)
    doc.add_paragraph()

    # --- 5.1 Gmail Integration ---
    doc.add_heading('5.1 Gmail Integration', level=2)

    doc.add_heading('Prerequisites', level=3)
    steps = [
        'A Google account with Gmail (this will be the monitored/protected mailbox).',
        'Access to Google Cloud Console (https://console.cloud.google.com).',
        'The Gmail API must be enabled in the Google Cloud project.',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph()

    doc.add_heading('Step 1: Create OAuth Credentials in Google Cloud Console', level=3)
    gcloud_steps = [
        'Go to https://console.cloud.google.com and select or create a project.',
        'Navigate to "APIs & Services" > "Library". Search for "Gmail API" and click "Enable".',
        'Go to "APIs & Services" > "Credentials". Click "Create Credentials" > "OAuth client ID".',
        'If prompted, configure the OAuth consent screen: set User Type to "External", enter "Stratos BEP" as the app name, and add your Gmail address as a test user.',
        'For Application Type, select "Web application" (NOT "Desktop application").',
        'Under "Authorized redirect URIs", add the exact URI shown on the Settings page. It looks like: http://YOUR_SERVER_IP/settings/gmail/callback/',
        'Click "Create". Download the JSON file when prompted.',
    ]
    for i, s in enumerate(gcloud_steps, 1):
        p = doc.add_paragraph()
        r = p.add_run(f'{i}. ')
        r.bold = True
        p.add_run(s)
    doc.add_paragraph()

    # Warning box
    warn = doc.add_paragraph()
    r = warn.add_run('IMPORTANT: ')
    r.bold = True
    r.font.color.rgb = RGBColor(220, 38, 38)
    warn.add_run('You must select "Web application" as the OAuth client type, not "Desktop application". '
                 'Desktop credentials will not work with the browser-based OAuth flow. '
                 'The redirect URI must match exactly — including http vs https and the trailing slash.')
    doc.add_paragraph()

    doc.add_heading('Step 2: Upload Credentials in Stratos', level=3)
    upload_steps = [
        'In the Settings page, find the "Gmail Integration" section.',
        'Click the "Choose File" button and select the downloaded JSON file.',
        'Click "Upload". You should see a green success message: "Gmail credentials uploaded."',
        'The status badge will remain "DISCONNECTED" until you complete Step 3.',
    ]
    for i, s in enumerate(upload_steps, 1):
        p = doc.add_paragraph()
        r = p.add_run(f'{i}. ')
        r.bold = True
        p.add_run(s)
    doc.add_paragraph()

    doc.add_heading('Step 3: Connect Gmail Account', level=3)
    connect_steps = [
        'Click the "Connect Gmail Account" button (blue button with envelope icon).',
        'Your browser will redirect to Google\'s login page.',
        'Sign in with the Gmail account you want Stratos to monitor.',
        'Review the permissions and click "Allow" to grant Stratos access.',
        'You will be redirected back to the Settings page.',
        'The status badge should now show "CONNECTED" in green, along with the connected email address.',
        'Stratos will begin fetching emails from this mailbox automatically.',
    ]
    for i, s in enumerate(connect_steps, 1):
        p = doc.add_paragraph()
        r = p.add_run(f'{i}. ')
        r.bold = True
        p.add_run(s)
    doc.add_paragraph()

    doc.add_heading('Disconnecting Gmail', level=3)
    doc.add_paragraph('To stop monitoring a Gmail account, click the red "Disconnect Gmail" button. '
                      'This removes the stored authentication token. Email fetching stops immediately. '
                      'You can reconnect at any time by repeating Step 3.')
    doc.add_paragraph()

    doc.add_heading('Troubleshooting Gmail Connection', level=3)
    trouble_table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
    trouble_table.rows[0].cells[0].text = 'Problem'
    trouble_table.rows[0].cells[1].text = 'Solution'
    troubles = [
        ('"Invalid credentials file" error', 'Ensure you downloaded a "Web application" type credential, not "Desktop application".'),
        ('"Redirect URI mismatch" from Google', 'The redirect URI in Google Console must match exactly: http://YOUR_IP/settings/gmail/callback/ (check http vs https and trailing slash).'),
        ('Status shows "EXPIRED"', 'The token has expired. Click "Disconnect", then "Connect Gmail Account" again to reauthorize.'),
        ('No emails being fetched', 'Check that Celery workers are running. Verify the fetch interval in Detection Settings. Check Celery logs on the server.'),
    ]
    for i, (prob, sol) in enumerate(troubles, 1):
        trouble_table.rows[i].cells[0].text = prob
        trouble_table.rows[i].cells[1].text = sol
    doc.add_paragraph()

    # --- 5.2 API Keys ---
    doc.add_heading('5.2 Threat Intelligence API Keys', level=2)

    doc.add_paragraph('Stratos integrates with two threat intelligence services for URL and IP reputation checking. '
                      'API keys are encrypted at rest using Fernet (AES-128-CBC) and are never stored as plaintext in the database.')
    doc.add_paragraph()

    doc.add_heading('VirusTotal API Key', level=3)
    vt_steps = [
        'Create a free account at https://www.virustotal.com/gui/join-us',
        'After registration, go to your profile and copy the API key.',
        'In Stratos Settings, paste the key into the "VirusTotal API Key" field.',
        'Click "Test" to verify the key works. You should see a green "VirusTotal API key is valid" message.',
        'Click "Save API Keys".',
    ]
    for i, s in enumerate(vt_steps, 1):
        p = doc.add_paragraph()
        r = p.add_run(f'{i}. ')
        r.bold = True
        p.add_run(s)

    p = doc.add_paragraph()
    r = p.add_run('Rate limit: ')
    r.bold = True
    p.add_run('4 requests per minute on the free tier. Stratos respects this limit automatically.')
    doc.add_paragraph()

    doc.add_heading('AbuseIPDB API Key', level=3)
    abuse_steps = [
        'Create a free account at https://www.abuseipdb.com/register',
        'Go to your account dashboard and generate an API key.',
        'In Stratos Settings, paste the key into the "AbuseIPDB API Key" field.',
        'Click "Test" to verify. You should see "AbuseIPDB API key is valid".',
        'Click "Save API Keys".',
    ]
    for i, s in enumerate(abuse_steps, 1):
        p = doc.add_paragraph()
        r = p.add_run(f'{i}. ')
        r.bold = True
        p.add_run(s)

    p = doc.add_paragraph()
    r = p.add_run('Rate limit: ')
    r.bold = True
    p.add_run('1,000 checks per day on the free tier.')
    doc.add_paragraph()

    doc.add_heading('What Happens Without API Keys', level=3)
    doc.add_paragraph('Stratos degrades gracefully when API keys are not configured:')
    degrade_table = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
    degrade_table.rows[0].cells[0].text = 'Missing Key'
    degrade_table.rows[0].cells[1].text = 'Impact'
    degrades = [
        ('No VirusTotal key', 'URL checking uses the local URLhaus database only. Most malicious URLs are still detected.'),
        ('No AbuseIPDB key', 'IP reputation checking is skipped. Other checks (keywords, headers, attachments) still run.'),
        ('No keys at all', 'System works using keyword analysis, email header authentication (SPF/DKIM/DMARC), attachment inspection, and the local threat intelligence database.'),
    ]
    for i, (key, impact) in enumerate(degrades, 1):
        degrade_table.rows[i].cells[0].text = key
        degrade_table.rows[i].cells[1].text = impact
    doc.add_paragraph()

    doc.add_heading('Key Security', level=3)
    security_points = [
        'API keys are encrypted using Fernet (AES-128-CBC) derived from the Django SECRET_KEY.',
        'Keys are displayed masked in the UI (e.g., "test****2345") — the full key is never shown.',
        'Only ADMIN users can view or modify API keys.',
        'Leave a field blank when saving to keep the existing key unchanged.',
        'Do NOT change the Django SECRET_KEY after saving API keys — it will make stored keys unreadable.',
    ]
    for s in security_points:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph()

    # --- 5.3 Detection Thresholds ---
    doc.add_heading('5.3 Detection Thresholds', level=2)

    doc.add_paragraph('Every email analyzed by Stratos receives a threat score from 0 to 100. '
                      'The score determines the verdict:')
    doc.add_paragraph()

    score_table = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
    score_table.rows[0].cells[0].text = 'Score Range'
    score_table.rows[0].cells[1].text = 'Verdict'
    score_table.rows[0].cells[2].text = 'Action'
    score_data = [
        ('Below Clean Threshold (default: < 25)', 'CLEAN', 'Email delivered normally'),
        ('Between thresholds (default: 25-69)', 'SUSPICIOUS', 'Email quarantined for review'),
        ('At/above Malicious Threshold (default: >= 70)', 'MALICIOUS', 'Email blocked automatically'),
    ]
    for i, (rng, verdict, action) in enumerate(score_data, 1):
        score_table.rows[i].cells[0].text = rng
        score_table.rows[i].cells[1].text = verdict
        score_table.rows[i].cells[2].text = action
    doc.add_paragraph()

    doc.add_heading('Adjusting Thresholds', level=3)
    threshold_steps = [
        'Use the "Clean Threshold" slider to set the upper bound of CLEAN emails (default: 25).',
        'Use the "Malicious Threshold" slider to set the lower bound of MALICIOUS emails (default: 70).',
        'The range between these two values becomes the SUSPICIOUS zone.',
        'Click "Save Detection Settings" to apply.',
    ]
    for i, s in enumerate(threshold_steps, 1):
        p = doc.add_paragraph()
        r = p.add_run(f'{i}. ')
        r.bold = True
        p.add_run(s)
    doc.add_paragraph()

    doc.add_heading('Recommended Threshold Configurations', level=3)
    rec_table = doc.add_table(rows=4, cols=4, style='Light Grid Accent 1')
    rec_table.rows[0].cells[0].text = 'Scenario'
    rec_table.rows[0].cells[1].text = 'Clean'
    rec_table.rows[0].cells[2].text = 'Malicious'
    rec_table.rows[0].cells[3].text = 'Notes'
    recs = [
        ('Default (balanced)', '25', '70', 'Good for most environments. Balanced false positives vs detection.'),
        ('High security', '15', '50', 'More emails quarantined. Higher false positive rate. Use when security is critical.'),
        ('Low noise', '35', '85', 'Fewer alerts. Risk of missing some threats. Use for low-risk environments.'),
    ]
    for i, (scenario, clean, mal, notes) in enumerate(recs, 1):
        rec_table.rows[i].cells[0].text = scenario
        rec_table.rows[i].cells[1].text = clean
        rec_table.rows[i].cells[2].text = mal
        rec_table.rows[i].cells[3].text = notes
    doc.add_paragraph()

    # --- 5.4 Fetch Interval & TI Sync ---
    doc.add_heading('5.4 Fetch Interval and TI Sync', level=2)

    doc.add_heading('Email Fetch Interval', level=3)
    doc.add_paragraph('Controls how frequently Stratos checks the connected Gmail mailbox for new emails.')
    fetch_points = [
        'Default: 10 seconds. Good for demos and viva presentations.',
        'Minimum: 5 seconds. Do not set lower to avoid Gmail API rate limiting.',
        'For production with high email volume: 30-60 seconds is recommended.',
        'Change the value in the "Fetch Interval" field and click "Save Detection Settings".',
    ]
    for s in fetch_points:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph()

    doc.add_heading('TI Feed Sync Toggle', level=3)
    doc.add_paragraph('The "Enable TI Feed Sync" checkbox controls whether Stratos automatically syncs '
                      'threat intelligence data from MalwareBazaar and URLhaus daily at 2:00 AM UTC.')
    ti_points = [
        'When enabled (default): New malware hashes and malicious domains are imported daily.',
        'When disabled: The existing TI database is still used for checks, but no new data is imported.',
        'You can always trigger a manual sync from the Threat Intel page (Admin > Threat Intel > Sync Now).',
        'Disabling sync does NOT delete existing TI data.',
    ]
    for s in ti_points:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph()

    # --- 5.5 DO and DO NOT ---
    doc.add_heading('5.5 Settings — DO and DO NOT', level=2)

    doc.add_heading('DO', level=3)
    dos = [
        'Use a "Web Application" OAuth client type in Google Cloud Console.',
        'Set the redirect URI to match your server URL exactly (including http/https and trailing slash).',
        'Test API keys after entering them using the "Test" button.',
        'Keep thresholds at defaults (25/70) unless you have a specific reason to change them.',
        'Monitor the dashboard for unusual detection patterns after changing thresholds.',
        'Back up your .env file before making changes to the Django SECRET_KEY.',
    ]
    for s in dos:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('DO NOT', level=3)
    donts = [
        'Share API keys or OAuth credentials with anyone.',
        'Set the clean threshold higher than the malicious threshold.',
        'Set the fetch interval below 5 seconds (Gmail rate limiting risk).',
        'Disconnect Gmail during an active email analysis.',
        'Upload a "Desktop Application" credentials file — it will not work for web OAuth.',
        'Change the Django SECRET_KEY after deployment — this will break encrypted API keys.',
        'Give the ADMIN role to users who do not need it.',
        'Disable TI sync without a reason — it reduces detection capability.',
    ]
    for s in donts:
        doc.add_paragraph(s, style='List Bullet')

    # Save
    doc.save(OUTPUT_PATH)
    print(f'\nPlaybook saved to: {OUTPUT_PATH}')


def main():
    parser = argparse.ArgumentParser(description='Generate Stratos BEP UI Playbook')
    parser.add_argument('--url', default='http://178.104.202.193', help='Base URL of the Stratos instance')
    args = parser.parse_args()

    base_url = args.url.rstrip('/')
    os.makedirs(SCREENSHOT_DIR, exist_ok=True)

    print(f'Generating playbook from {base_url}')

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(viewport={'width': 1440, 'height': 900})
        page = context.new_page()

        admin_shots = capture_admin_screenshots(page, base_url)
        analyst_shots = capture_analyst_screenshots(page, base_url)
        viewer_shots = capture_viewer_screenshots(page, base_url)

        browser.close()

    build_docx(admin_shots, analyst_shots, viewer_shots)


if __name__ == '__main__':
    main()
